"""Extract plain text from uploaded PDF, DOCX, or TXT files."""
import io
import logging

log = logging.getLogger(__name__)


def extract_text(content: bytes, filename: str) -> str:
    """Return all readable text from a file's raw bytes."""
    fname = (filename or "").lower()

    if fname.endswith(".pdf"):
        return _extract_pdf(content)
    if fname.endswith(".docx"):
        return _extract_docx(content)
    if fname.endswith(".doc"):
        return _extract_doc_fallback(content)
    # Plain text / CSV / TXT
    try:
        return content.decode("utf-8", errors="replace")
    except Exception:
        return ""


def _extract_pdf(content: bytes) -> str:
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(stream=content, filetype="pdf")
        pages = []
        for page in doc:
            pages.append(page.get_text())
        doc.close()
        return "\n".join(pages)
    except Exception as e:
        log.warning("PDF extraction failed: %s", e)
        return ""


def _extract_docx(content: bytes) -> str:
    try:
        from docx import Document
        doc = Document(io.BytesIO(content))
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        # Also extract table cells
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    parts.append(row_text)
        return "\n".join(parts)
    except Exception as e:
        log.warning("DOCX extraction failed: %s", e)
        return ""


def _extract_doc_fallback(content: bytes) -> str:
    """Best-effort text extraction from legacy .doc files."""
    try:
        # Try treating as raw binary and pulling ASCII text runs
        text = content.decode("latin-1", errors="replace")
        printable = "".join(c if c.isprintable() or c in "\n\r\t" else " " for c in text)
        # Collapse whitespace
        import re
        cleaned = re.sub(r" {3,}", " ", printable)
        cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
        return cleaned[:50000]  # cap at 50k chars for old .doc
    except Exception:
        return ""
